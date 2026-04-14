"""
Task: S002/T002 — DocxParser tests (TDD)
Spec: docs/document-parser/spec/document-parser.spec.md
"""
import io
import time

import pytest
from docx import Document as DocxDocument  # type: ignore[import]

from backend.rag.parser.docx_parser import DocxParser
from backend.rag.parser.base import ParseError


# ---------------------------------------------------------------------------
# Helpers — minimal DOCX byte builders
# ---------------------------------------------------------------------------

def _make_docx_bytes(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    """Build a minimal DOCX with given paragraphs and an optional table."""
    doc = DocxDocument()
    for text in paragraphs:
        doc.add_paragraph(text)
    if table_rows:
        cols = max(len(row) for row in table_rows)
        tbl = doc.add_table(rows=len(table_rows), cols=cols)
        for r_idx, row in enumerate(table_rows):
            for c_idx, cell_text in enumerate(row):
                tbl.rows[r_idx].cells[c_idx].text = cell_text
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_empty_docx_bytes() -> bytes:
    """Build a DOCX with no paragraphs and no tables."""
    doc = DocxDocument()
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

def test_paragraphs_and_table_both_extracted():
    """Multi-section DOCX with paragraphs + table → both included in output."""
    data = _make_docx_bytes(
        paragraphs=["Introduction paragraph", "Second paragraph"],
        table_rows=[["Cell A1", "Cell A2"], ["Cell B1", "Cell B2"]],
    )
    parser = DocxParser()
    doc = parser.parse(data)

    assert "Introduction paragraph" in doc.text
    assert "Second paragraph" in doc.text
    assert "Cell A1" in doc.text
    assert "Cell B2" in doc.text
    assert doc.lang is None
    assert doc.metadata["source_format"] == "docx"


def test_table_cell_and_row_joining():
    """Table cells joined with ' | ', rows joined with newline."""
    data = _make_docx_bytes(
        paragraphs=["Intro"],
        table_rows=[["Alpha", "Beta", "Gamma"], ["Delta", "Epsilon", "Zeta"]],
    )
    parser = DocxParser()
    doc = parser.parse(data)

    assert "Alpha | Beta | Gamma" in doc.text
    assert "Delta | Epsilon | Zeta" in doc.text


def test_empty_docx_raises_err_parse_failed():
    """Empty DOCX (no paragraphs, no tables) → ParseError(ERR_PARSE_FAILED)."""
    data = _make_empty_docx_bytes()
    parser = DocxParser()

    with pytest.raises(ParseError) as exc_info:
        parser.parse(data)
    assert exc_info.value.code == "ERR_PARSE_FAILED"


def test_corrupt_docx_raises_err_parse_failed():
    """Corrupt bytes → ParseError(ERR_PARSE_FAILED)."""
    parser = DocxParser()

    with pytest.raises(ParseError) as exc_info:
        parser.parse(b"PK CORRUPT GARBAGE NOT A DOCX ####")
    assert exc_info.value.code == "ERR_PARSE_FAILED"


def test_section_count_in_metadata():
    """section_count = number of non-empty paragraphs."""
    data = _make_docx_bytes(paragraphs=["Para one", "Para two", "Para three"])
    parser = DocxParser()
    doc = parser.parse(data)

    assert doc.metadata["section_count"] == 3
    assert doc.metadata["source_format"] == "docx"
    assert doc.lang is None


@pytest.mark.slow
def test_perf_large_docx():
    """≤200KB DOCX should parse in under 1 second."""
    paragraphs = [f"This is paragraph number {i} in the performance test." for i in range(1, 201)]
    data = _make_docx_bytes(paragraphs=paragraphs)
    assert len(data) <= 200 * 1024, "Fixture unexpectedly large"

    parser = DocxParser()
    start = time.monotonic()
    doc = parser.parse(data)
    elapsed = time.monotonic() - start

    assert doc.metadata["section_count"] == 200
    assert elapsed < 1.0, f"Parsing took {elapsed:.2f}s (limit: 1s)"
