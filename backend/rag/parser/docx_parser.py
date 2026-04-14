# Spec: docs/document-parser/spec/document-parser.spec.md
# Task: S002/T002 — DocxParser
# Decision: D01 (python-docx)
import io

from docx import Document  # type: ignore[import]

from .base import ParseError, ParsedDocument, ParserBase


class DocxParser(ParserBase):
    def parse(self, data: bytes) -> ParsedDocument:
        try:
            doc = Document(io.BytesIO(data))
        except Exception as e:
            raise ParseError("ERR_PARSE_FAILED", str(e)) from e

        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))

        if not parts:
            raise ParseError("ERR_PARSE_FAILED", "Document contains no text content")

        section_count = sum(1 for p in doc.paragraphs if p.text.strip())
        return ParsedDocument(
            text="\n".join(parts),
            lang=None,
            metadata={
                "section_count": section_count,
                "source_format": "docx",
            },
        )
